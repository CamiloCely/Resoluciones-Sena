import os
import re
import time
import datetime
import pandas as pd
from pypdf import PdfReader
from docx import Document
import streamlit as st

st.set_page_config(
    page_title="Generador de Resoluciones SENA",
    page_icon="🏛️",
    layout="centered"
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# --- ADMINISTRACIÓN DE BASES DE DATOS ---
st.sidebar.title("⚙️ Administración de Bases de Datos")

nuevo_excel = st.sidebar.file_uploader("📊 Actualizar Excel Kactus / Vacaciones", type=["xlsx", "xls"])
if nuevo_excel is not None:
    path_nuevo_excel = os.path.join(BASE_DIR, "Kactus_Actualizado.xlsx")
    with open(path_nuevo_excel, "wb") as f:
        f.write(nuevo_excel.getbuffer())
    st.sidebar.success("✅ Base de Vacaciones actualizada.")

nuevo_maestro = st.sidebar.file_uploader("📋 Actualizar Maestro por Dependencias (PDF)", type=["pdf"])
if nuevo_maestro is not None:
    path_nuevo_maestro = os.path.join(BASE_DIR, "MAESTRO_CARGOS_ACTUALIZADO.pdf")
    with open(path_nuevo_maestro, "wb") as f:
        f.write(nuevo_maestro.getbuffer())
    st.sidebar.success("✅ Maestro de Dependencias actualizado.")

st.sidebar.divider()

archivos_excel = [f for f in os.listdir(BASE_DIR) if f.lower().endswith(('.xlsx', '.xls'))]
EXCEL_HISTORIAL = os.path.join(BASE_DIR, "Kactus_Actualizado.xlsx") if os.path.exists(os.path.join(BASE_DIR, "Kactus_Actualizado.xlsx")) else (os.path.join(BASE_DIR, archivos_excel[0]) if archivos_excel else None)

archivos_word = [f for f in os.listdir(BASE_DIR) if f.lower().endswith('.docx') and not f.startswith('~$')]
PLANTILLA_WORD = os.path.join(BASE_DIR, archivos_word[0]) if archivos_word else None

archivos_pdf_maestro = [f for f in os.listdir(BASE_DIR) if "MAESTRO" in f.upper() and f.lower().endswith('.pdf')]
MAESTRO_CARGOS = os.path.join(BASE_DIR, "MAESTRO_CARGOS_ACTUALIZADO.pdf") if os.path.exists(os.path.join(BASE_DIR, "MAESTRO_CARGOS_ACTUALIZADO.pdf")) else (os.path.join(BASE_DIR, archivos_pdf_maestro[0]) if archivos_pdf_maestro else None)

if st.sidebar.button("🔄 Reiniciar Memoria / Forzar Limpieza"):
    st.cache_data.clear()
    st.rerun()

FESTIVOS_COLOMBIA = [
    datetime.date(2026, 1, 1),   datetime.date(2026, 1, 12),  datetime.date(2026, 3, 23),
    datetime.date(2026, 4, 2),   datetime.date(2026, 4, 3),   datetime.date(2026, 5, 1),
    datetime.date(2026, 5, 18),  datetime.date(2026, 6, 8),   datetime.date(2026, 6, 15),
    datetime.date(2026, 6, 29),  datetime.date(2026, 7, 20),  datetime.date(2026, 8, 7),
    datetime.date(2026, 10, 12), datetime.date(2026, 11, 2),  datetime.date(2026, 11, 16),
    datetime.date(2026, 12, 8),  datetime.date(2026, 12, 25),
]

def calcular_fecha_fin(fecha_inicio, dias_habiles=15):
    fecha_actual = fecha_inicio
    dias_contados = 0
    while dias_contados < dias_habiles:
        if fecha_actual.weekday() < 5 and fecha_actual not in FESTIVOS_COLOMBIA:
            dias_contados += 1
        if dias_contados < dias_habiles:
            fecha_actual += datetime.timedelta(days=1)
    return fecha_actual

def extraer_datos_pdf(file_bytes, filename_pdf=""):
    lector = PdfReader(file_bytes)
    texto = ""
    for pag in lector.pages:
        txt = pag.extract_text()
        if txt:
            texto += txt + "\n"
            
    meses_dict = {1:"enero", 2:"febrero", 3:"marzo", 4:"abril", 5:"mayo", 6:"junio", 7:"julio", 8:"agosto", 9:"septiembre", 10:"octubre", 11:"noviembre", 12:"diciembre"}
    meses_nom = {"enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6, "julio": 7, "agosto": 8, "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12}

    texto_unificado = " ".join(texto.split())

    # 1. RADICADO
    rad_match = re.search(r"(\d{2}\-\d{1,2}\-\d{4}\-\d{4,8})", texto_unificado)
    if not rad_match:
        rad_match = re.search(r"No:\s*([\d\-]{10,25})", texto_unificado, re.IGNORECASE)
    if not rad_match and filename_pdf:
        rad_match = re.search(r"(\d{2}\-\d{1,2}\-\d{4}\-\d{4,8})", filename_pdf)

    radicado = rad_match.group(1).strip() if rad_match else ""

    # 2. FECHA DE RADICACIÓN MULTINIVEL
    fecha_rad_str = ""
    match_fecha_txt = re.search(r"(\d{1,2})\s+de\s+([a-zA-Z]+)\s+de\s+(\d{4})", texto_unificado, re.IGNORECASE)
    if match_fecha_txt:
        d_txt = int(match_fecha_txt.group(1))
        m_txt = match_fecha_txt.group(2).lower()
        a_txt = match_fecha_txt.group(3)
        fecha_rad_str = f"{d_txt:02d} de {m_txt} de {a_txt}"

    if not fecha_rad_str:
        sticker_match = re.search(r"(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})", texto_unificado)
        if sticker_match:
            dia_s = int(sticker_match.group(1))
            mes_s = int(sticker_match.group(2))
            ano_s = sticker_match.group(3)
            fecha_rad_str = f"{dia_s:02d} de {meses_dict.get(mes_s, 'junio')} de {ano_s}"

    if not fecha_rad_str:
        hoy_temp = datetime.date.today()
        fecha_rad_str = f"{hoy_temp.day:02d} de {meses_dict.get(hoy_temp.month, 'junio')} de {hoy_temp.year}"

    # 3. CÉDULAS
    todas_cedulas = re.findall(r"(\d{7,10})", texto_unificado)

    # 4. PERÍODO CAUSADO
    p_inicio, p_fin = "", ""
    periodo_match = re.search(r"(?:comprendido\s+entre\s+el|periodo\s+del|periodo\s+causado\s+del)\s+(\d{1,2}\s+de\s+[a-zA-Z]+\s+de\s+\d{4})\s+(?:al|hasta\s+el|y\s+el)\s+(\d{1,2}\s+de\s+[a-zA-Z]+\s+de\s+\d{4})", texto_unificado, re.IGNORECASE)
    if periodo_match:
        p_inicio = periodo_match.group(1).strip()
        p_fin = periodo_match.group(2).strip()
    else:
        patron_fechas = re.findall(r"(\d{1,2}\s+de\s+[a-zA-Z]+\s+de\s+\d{4})", texto_unificado, re.IGNORECASE)
        if len(patron_fechas) >= 2:
            p_inicio = patron_fechas[0]
            p_fin = patron_fechas[1]

    # 5. FECHA DE DISFRUTE
    disfrute_match = re.search(r"a\s+partir\s+del\s+(\d{1,2}\s+de\s+[a-zA-Z]+(?:\s+de\s+\d{4})?)", texto_unificado, re.IGNORECASE)
    fecha_disfrute_obj = datetime.date.today()
    if disfrute_match:
        raw_d = disfrute_match.group(1).lower().replace(".", "").strip()
        partes = raw_d.split()
        if len(partes) >= 3:
            d_dia = int(partes[0])
            d_mes = meses_nom.get(partes[2], 7)
            d_ano = int(partes[4]) if len(partes) >= 5 else 2026
            fecha_disfrute_obj = datetime.date(d_ano, d_mes, d_dia)

    return {
        "radicado": radicado,
        "fecha_radicado": fecha_rad_str,
        "periodo_inicio": p_inicio,
        "periodo_fin": p_fin,
        "fecha_inicio_obj": fecha_disfrute_obj,
        "cedulas_extraidas": todas_cedulas,
        "texto_completo": texto_unificado.upper()
    }

def obtener_datos_centro_y_firmante(codigo_dep):
    DATOS_CENTROS = {
        "9110": {
            "centro": "Centro de Desarrollo Agropecuario y Agroindustrial de la regional Boyacá",
            "titulo_encabezado": "SUBDIRECTORA (E) DEL CENTRO DE DESARROLLO AGROPECUARIO Y AGROINDUSTRIAL DEL SERVICIO NACIONAL DE APRENDIZAJE \"SENA\" REGIONAL BOYACÁ",
            "ciudad": "Duitama",
            "jefe_nombre": "Enith Yadira Ramírez Camargo",
            "jefe_cargo": "Subdirectora de Centro (E)"
        },
        "9111": {
            "centro": "Centro Minero de la regional Boyacá",
            "titulo_encabezado": "SUBDIRECTORA (E) DEL CENTRO MINERO DEL SERVICIO NACIONAL DE APRENDIZAJE \"SENA\" REGIONAL BOYACÁ",
            "ciudad": "Sogamoso",
            "jefe_nombre": "Angela María Montoya Castro",
            "jefe_cargo": "Subdirectora (E) Centro Minero Regional Boyacá"
        },
        "9305": {
            "centro": "Centro de Gestión Administrativa y Fortalecimiento Empresarial de la regional Boyacá",
            "titulo_encabezado": "SUBDIRECTOR (E) DEL CENTRO DE GESTIÓN ADMINISTRATIVA Y FORTALECIMIENTO EMPRESARIAL DEL SERVICIO NACIONAL DE APRENDIZAJE \"SENA\" REGIONAL BOYACÁ",
            "ciudad": "Tunja",
            "jefe_nombre": "Subdirector CGAFE",
            "jefe_cargo": "Subdirector de Centro (E)"
        },
        "9514": {
            "centro": "Centro Industrial de Mantenimiento y Manufactura de la regional Boyacá",
            "titulo_encabezado": "SUBDIRECTOR (E) DEL CENTRO INDUSTRIAL DE MANTENIMIENTO Y MANUFACTURA DEL SERVICIO NACIONAL DE APRENDIZAJE \"SENA\" REGIONAL BOYACÁ",
            "ciudad": "Sogamoso",
            "jefe_nombre": "Consuelo Alexandra Barrera Coronado",
            "jefe_cargo": "Subdirectora (E) CIMM"
        },
        "1010": {
            "centro": "Despacho Dirección Regional Boyacá",
            "titulo_encabezado": "DIRECTOR REGIONAL DEL SERVICIO NACIONAL DE APRENDIZAJE \"SENA\" REGIONAL BOYACÁ",
            "ciudad": "Tunja",
            "jefe_nombre": "Director Regional",
            "jefe_cargo": "Director Regional Boyacá"
        }
    }
    return DATOS_CENTROS.get(str(codigo_dep), DATOS_CENTROS["1010"])

def obtener_cargo_y_dep(nombre_empleado, cedula=None):
    cargo_oficial = "Profesional G04"
    codigo_dep = "1010"

    if not MAESTRO_CARGOS or not os.path.exists(MAESTRO_CARGOS):
        return cargo_oficial, codigo_dep

    lector = PdfReader(MAESTRO_CARGOS)
    nombre_buscar = nombre_empleado.upper().strip()

    for pag in lector.pages:
        lineas = pag.extract_text().split("\n")
        for linea in lineas:
            if "DEPENDENCIA:" in linea:
                for cod in ["9110", "9111", "9305", "9514", "1010"]:
                    if cod in linea:
                        codigo_dep = cod
            
            coincide_cedula = cedula and cedula in linea
            partes_nom = nombre_buscar.split()
            coincide_nombre = len(partes_nom) >= 2 and partes_nom[0] in linea.upper() and partes_nom[-1] in linea.upper()
            
            if coincide_cedula or coincide_nombre:
                match_cargo = re.search(r"(Instructor\s+G\d+|Profesional\s+G\d+(?:\s*\(e\))?|Tecnico\s+G\d+|Secretaria\s+G\d+|Auxiliar\s+G\d+|Subdirector\s+De\s+Centro|Oficial\s+Mantto[^\d]*G\d+)", linea, re.IGNORECASE)
                if match_cargo:
                    cargo_oficial = match_cargo.group(1).strip()
                return cargo_oficial, codigo_dep
                
    return cargo_oficial, codigo_dep

def reemplazar_respetando_formato(doc, dic_reemplazos):
    def procesar_p(p):
        for k, v in dic_reemplazos.items():
            if k in p.text:
                full_text = p.text.replace(k, str(v))
                runs_no_img = [r for r in p.runs if not any(tag in r._element.xml for tag in ['w:drawing', 'w:pict', 'a:blip', 'v:shape'])]
                if runs_no_img:
                    runs_no_img[0].text = full_text
                    for r in runs_no_img[1:]:
                        r.text = ""

    for p in doc.paragraphs:
        procesar_p(p)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    procesar_p(p)

# --- INTERFAZ STREAMLIT ---
st.title("🏛️ Sistema Automático de Resoluciones de Vacaciones")
st.markdown("Carga la carta de solicitud enviada por el funcionario (PDF) para generar la resolución oficial en Word.")

if not EXCEL_HISTORIAL or not PLANTILLA_WORD:
    st.error("⚠️ Verifica que la plantilla .docx y la base Excel estén configuradas.")
else:
    archivo_pdf = st.file_uploader("Arrastra aquí la carta de solicitud recibida (.pdf)", type=["pdf"], key="pdf_uploader")

    if archivo_pdf is not None:
        datos_carta = extraer_datos_pdf(archivo_pdf, archivo_pdf.name)
        
        xls = pd.ExcelFile(EXCEL_HISTORIAL)
        nombre_hoja = 'KactuS - KNmVacac' if 'KactuS - KNmVacac' in xls.sheet_names else xls.sheet_names[0]
        df_kactus = pd.read_excel(EXCEL_HISTORIAL, sheet_name=nombre_hoja)

        nombres_limpios = df_kactus['Nombre del Empleado'].fillna('').astype(str)
        apellidos_limpios = df_kactus['Apellidos Empleado'].fillna('').astype(str)
        df_kactus['Nombre_Completo'] = (nombres_limpios + " " + apellidos_limpios).str.strip().str.upper()
        
        lista_funcionarios = sorted([n for n in df_kactus['Nombre_Completo'].unique() if len(n) > 2])

        # BÚSQUEDA DEL SOLICITANTE POR CÉDULA
        indice_sugerido = 0
        encontrado = False

        for c_ext in datos_carta['cedulas_extraidas']:
            filas_c = df_kactus[df_kactus['Identificación'].astype(str).str.contains(c_ext)]
            if not filas_c.empty:
                nom_c = filas_c.iloc[0]['Nombre_Completo']
                if nom_c in lista_funcionarios:
                    indice_sugerido = lista_funcionarios.index(nom_c)
                    encontrado = True
                    break

        if not encontrado:
            for idx, nom in enumerate(lista_funcionarios):
                partes = nom.split()
                if len(partes) >= 2:
                    if partes[0] in datos_carta['texto_completo'] and partes[-1] in datos_carta['texto_completo']:
                        if "ENITH" not in nom and "NIDIA" not in nom and "NORA" not in nom:
                            indice_sugerido = idx
                            break

        st.success("📄 Carta analizada correctamente.")
        st.markdown("### 👤 Confirmación del Solicitante:")
        
        solicitante_elegido = st.selectbox(
            "Verifica o selecciona el funcionario que solicita las vacaciones:",
            options=lista_funcionarios,
            index=indice_sugerido
        )

        fila_encontrada = df_kactus[df_kactus['Nombre_Completo'] == solicitante_elegido].iloc[0]
        
        cedula_num = int(fila_encontrada['Identificación'])
        cedula_puntos = f"{cedula_num:,}".replace(",", ".")
        nombre_completo = solicitante_elegido

        st.markdown("---")
        st.markdown("### 📅 Ajusta o confirma las fechas e información extraída de la carta:")
        col1, col2 = st.columns(2)
        with col1:
            val_rad = datos_carta['radicado'] if datos_carta['radicado'] else "[COMPLETAR RADICADO]"
            val_f_rad = datos_carta['fecha_radicado'] if datos_carta['fecha_radicado'] else "[COMPLETAR FECHA RADICADO]"
            radicado_final = st.text_input("Número de Radicado:", value=val_rad)
            fecha_rad_final = st.text_input("Fecha del Radicado:", value=val_f_rad)
            p_ini_final = st.text_input("Período Causado (Inicio):", value=datos_carta['periodo_inicio'])
        with col2:
            p_fin_final = st.text_input("Período Causado (Fin):", value=datos_carta['periodo_fin'])
            f_ini_obj_final = st.date_input("Fecha Inicio Disfrute:", value=datos_carta['fecha_inicio_obj'])

        st.markdown("---")

        if st.button("⚡ Generar Resolución en Word"):
            genero = str(fila_encontrada.get('Sexo', '')).upper()
            if 'F' in genero or nombre_completo.startswith(('BLANCA', 'MARIA', 'ANGELA', 'NEILA', 'NIDIA', 'YADIRA', 'KATHERINE', 'SANDRA', 'PATRICIA', 'LILIANA', 'CLAUDIA', 'SONIA', 'ROSA', 'ANA', 'CONSUELO', 'NORA', 'IRMA')):
                texto_funcionario = "la funcionaria"
                texto_funcionario_a = "a la funcionaria"
            else:
                texto_funcionario = "el funcionario"
                texto_funcionario_a = "al funcionario"

            cargo, cod_dep = obtener_cargo_y_dep(nombre_completo, str(cedula_num))
            info_centro = obtener_datos_centro_y_firmante(cod_dep)
            
            meses_esp = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
            
            fecha_fin_obj = calcular_fecha_fin(f_ini_obj_final, 15)
            
            dia_fin_str = f"{fecha_fin_obj.day:02d}" if fecha_fin_obj.day < 10 else f"{fecha_fin_obj.day}"
            fecha_fin_str = f"{dia_fin_str} de {meses_esp[fecha_fin_obj.month - 1]} de {fecha_fin_obj.year}"
            
            dia_ini_str = f"{f_ini_obj_final.day:02d}" if f_ini_obj_final.day < 10 else f"{f_ini_obj_final.day}"
            fecha_inicio_formateada = f"{dia_ini_str} de {meses_esp[f_ini_obj_final.month - 1]} de {f_ini_obj_final.year}"

            hoy = datetime.date.today()
            fecha_hoy_str = f"{hoy.day:02d} de {meses_esp[hoy.month - 1]} de {hoy.year}"

            doc = Document(PLANTILLA_WORD)
            
            reemplazos = {
                "[TITULO_DIRECTOR_COMPLETO]": info_centro["titulo_encabezado"],
                "[TEXTO_FUNCIONARIO]": texto_funcionario,
                "[TEXTO_FUNCIONARIO_A]": texto_funcionario_a,
                "[NOMBRE_EMPLEADO]": nombre_completo,
                "[CEDULA]": cedula_puntos,
                "[CARGO]": cargo,
                "[CENTRO_FORMACION]": info_centro["centro"],
                "[RADICADO]": radicado_final,
                "[FECHA_RADICADO]": fecha_rad_final,
                "[FECHA_INICIO]": fecha_inicio_formateada,
                "[FECHA_FIN]": fecha_fin_str,
                "[PERIODO_INICIO]": p_ini_final,
                "[PERIODO_FIN]": p_fin_final,
                "[CIUDAD_CENTRO]": info_centro["ciudad"],
                "[FECHA_HOY]": fecha_hoy_str,
                "[NOMBRE_JEFE_FIRMA]": info_centro["jefe_nombre"],
                "[CARGO_JEFE_FIRMA]": info_centro["jefe_cargo"]
            }
            
            reemplazar_respetando_formato(doc, reemplazos)

            timestamp_unico = int(time.time())
            nombre_archivo_salida = f"Resolucion_Vacaciones_{nombre_completo.replace(' ', '_')}_{timestamp_unico}.docx"
            salida_path = os.path.join(BASE_DIR, nombre_archivo_salida)
            doc.save(salida_path)

            st.balloons()
            st.success(f"✅ ¡Resolución generada con éxito!")

            with open(salida_path, "rb") as file_docx:
                st.download_button(
                    label=f"📥 DESCARGAR RESOLUCIÓN DE {nombre_completo}",
                    data=file_docx,
                    file_name=f"Resolucion_Vacaciones_{nombre_completo.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
